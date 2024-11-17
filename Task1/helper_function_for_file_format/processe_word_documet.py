from docx import Document

def process_word_document(file_path):
    """Extract information from a Word document."""
    try:
        doc = Document(file_path)

        # Extract title (assume the first paragraph is the title)
        title = doc.paragraphs[0].text.strip() if doc.paragraphs else ''

        # Extract all tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        # Extract plain text (all paragraphs)
        plain_text = '\n'.join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

        return {
            'title': title,
            'table': tables,
            'key_value_pairs': {},  # Add logic for key-value pairs if required
            'plain_text': plain_text
        }
    except Exception as e:
        print(f"Error processing Word document: {str(e)}")
        return {}
